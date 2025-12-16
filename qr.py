import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser
import qrcode
from qrcode.constants import ERROR_CORRECT_M
from PIL import Image, ImageTk
import os
from docx import Document
from docx.shared import Inches

class QRCodeGeneratorPro:
    def __init__(self, root):
        self.root = root
        self.root.title("QR Pro Generator – Professional QR Kod Studiyasi")
        self.root.geometry("900x720")

        self.fill_color = "#000000"
        self.back_color = "#FFFFFF"
        self.box_size = tk.IntVar(value=10)

        self.setup_ui()

    def setup_ui(self):
        main = ttk.Frame(self.root, padding=10)
        main.pack(fill=tk.BOTH, expand=True)

        # HEADER
        ttk.Label(
            main,
            text="🎯 QR Pro Generator",
            font=("Segoe UI", 20, "bold")
        ).pack(anchor=tk.W)

        # CONTENT
        content = ttk.Frame(main)
        content.pack(fill=tk.BOTH, expand=True, pady=10)

        left = ttk.Frame(content)
        left.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))

        right = ttk.LabelFrame(content, text="QR Kod Ko‘rinishi", padding=10)
        right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        self.create_input(left)
        self.create_settings(left)
        self.create_preview(right)

    # 🔹 TEXT INPUT
    def create_input(self, parent):
        frame = ttk.LabelFrame(parent, text="QR Kod Matni", padding=10)
        frame.pack(fill=tk.X)

        self.text_input = tk.Text(frame, height=6, wrap=tk.WORD)
        self.text_input.pack(fill=tk.X)
        self.text_input.insert("1.0", "https://misol-link.uz")

        ttk.Button(
            frame,
            text="🔄 QR Kod Yaratish",
            command=self.generate_qr
        ).pack(fill=tk.X, pady=5)

    # 🔹 SETTINGS
    def create_settings(self, parent):
        frame = ttk.LabelFrame(parent, text="Sozlamalar", padding=10)
        frame.pack(fill=tk.X, pady=10)

        ttk.Label(frame, text="Box size:").pack(anchor=tk.W)
        ttk.Scale(
            frame,
            from_=5,
            to=20,
            variable=self.box_size,
            orient=tk.HORIZONTAL
        ).pack(fill=tk.X)

        ttk.Button(
            frame,
            text="🎨 QR Rangini Tanlash",
            command=self.choose_color
        ).pack(fill=tk.X, pady=5)

    # 🔹 PREVIEW + SAVE BUTTONS
    def create_preview(self, parent):
        self.qr_label = ttk.Label(
            parent,
            text="QR kod bu yerda ko‘rinadi",
            anchor=tk.CENTER,
            background="white"
        )
        self.qr_label.pack(fill=tk.BOTH, expand=True)

        btns = ttk.Frame(parent)
        btns.pack(pady=5)

        ttk.Button(btns, text="🖼 PNG", command=lambda: self.save_qr("png")).pack(side=tk.LEFT, padx=5)
        ttk.Button(btns, text="📄 PDF", command=lambda: self.save_qr("pdf")).pack(side=tk.LEFT, padx=5)
        ttk.Button(btns, text="📝 Word", command=lambda: self.save_qr("docx")).pack(side=tk.LEFT, padx=5)

    # 🔹 FUNCTIONS
    def choose_color(self):
        color = colorchooser.askcolor()[1]
        if color:
            self.fill_color = color
            self.generate_qr()

    def generate_qr(self):
        text = self.text_input.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("Ogohlantirish", "Matn kiriting")
            return

        qr = qrcode.QRCode(
            error_correction=ERROR_CORRECT_M,
            box_size=self.box_size.get(),
            border=4
        )
        qr.add_data(text)
        qr.make(fit=True)

        img = qr.make_image(
            fill_color=self.fill_color,
            back_color=self.back_color
        ).convert("RGB")

        self.current_img = img

        preview = img.copy()
        preview.thumbnail((420, 420))
        self.qr_photo = ImageTk.PhotoImage(preview)
        self.qr_label.config(image=self.qr_photo, text="")

    def save_qr(self, fmt):
        if not hasattr(self, "current_img"):
            messagebox.showwarning("Ogohlantirish", "Avval QR kod yarating")
            return

        filetypes = {
            "png": [("PNG fayl", "*.png")],
            "pdf": [("PDF fayl", "*.pdf")],
            "docx": [("Word fayl", "*.docx")]
        }

        path = filedialog.asksaveasfilename(
            defaultextension=f".{fmt}",
            filetypes=filetypes[fmt]
        )

        if not path:
            return

        try:
            if fmt == "png":
                self.current_img.save(path, "PNG")

            elif fmt == "pdf":
                self.current_img.convert("RGB").save(path, "PDF", resolution=300)

            elif fmt == "docx":
                temp = "temp_qr.png"
                self.current_img.save(temp)

                doc = Document()
                doc.add_heading("QR Kod", level=1)
                doc.add_picture(temp, width=Inches(3))
                doc.save(path)

                os.remove(temp)

            messagebox.showinfo("Muvaffaqiyatli", f"{fmt.upper()} formatda saqlandi")

        except Exception as e:
            messagebox.showerror("Xatolik", str(e))


def main():
    root = tk.Tk()
    QRCodeGeneratorPro(root)
    root.mainloop()


if __name__ == "__main__":
    main()
